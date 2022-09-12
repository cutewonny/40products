import imp
from docx import Document
from openpyxl import load_workbook
import docx
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx2pdf import convert

load_wb = load_workbook(r'/Users/devwonny/Documents/40products/12. 엑셀의 정보를 불러와 수료증 자동 생성/엑셀생성.xlsx')
load_ws = load_wb.active

name_list = [] #1번 항목 : 이름
birthday_list = [] #2번 항목: 생일
ho_list = [] #3번 항목: 사원번호

print(load_ws.max_row) # 3개의 row
for i in range(1, load_ws.max_row +1):
    name_list.append(load_ws.cell(i,1).value)
    birthday_list.append(load_ws.cell(i,2).value)
    ho_list.append(load_ws.cell(i,3).value)

print(name_list)
print(birthday_list)
print(ho_list)

for i in range(len(name_list)):
    doc = docx.Document(r'/Users/devwonny/Documents/40products/12. 엑셀의 정보를 불러와 수료증 자동 생성/수료증양식.docx')

    style = doc.styles['Normal']
    style.font.name = '나눔고딕'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '나눔고딕')
    style.font.size = docx.shared.Pt(12)

    para = doc.add_paragraph()
    run = para.add_run('\n\n')
    run = para.add_run(f" 제  {ho_list[i]} 호\n") 
    run.font.name = '나눔고딕'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '나눔고딕')
    style.font.size = docx.shared.Pt(20)

    para = doc.add_paragraph()
    # run = para.add_run('\n\n')
    run = para.add_run('수 료 증')
    run.font.name = '나눔고딕'
    run.bold = True
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '나눔고딕')
    style.font.size = docx.shared.Pt(20)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    para = doc.add_paragraph()
    run = para.add_run('\n\n')
    run = para.add_run(f'        성       명: {name_list[i]}\n') 
    run.font.name = '나눔고딕'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '나눔고딕')
    run.font.size = docx.shared.Pt(20)
    run = para.add_run(f'        생 년 월 일: {birthday_list[i]}\n') 
    run.font.name = '나눔고딕'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '나눔고딕')
    run.font.size = docx.shared.Pt(20)
    run = para.add_run('        교 육 과 정: 파이썬과 40개의 작품들\n')
    run.font.name = '나눔고딕'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '나눔고딕')
    run.font.size = docx.shared.Pt(20) 
    run = para.add_run('        교 육 날 짜: 2021.08.05~2021.09.09\n') 
    run.font.name = '나눔고딕'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '나눔고딕')
    run.font.size = docx.shared.Pt(20)

    para = doc.add_paragraph() 
    run = para.add_run('\n\n')
    run = para.add_run('        위 사람은 파이썬과 40개의 작품들 교육과정을\n') 
    run.font.name = '나눔고딕'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '나눔고딕')
    run.font.size = docx.shared.Pt(20)
    run = para.add_run('        이수하였으므로 이 증서를 수여 합니다.\n') 
    run.font.name = '나눔고딕'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '나눔고딕')
    run.font.size = docx.shared.Pt(20)

    para = doc.add_paragraph()
    run = para.add_run('\n')
    run = para.add_run('2021.09.19') 
    run.font.name = '나눔고딕'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '나눔고딕')
    run.font.size = docx.shared.Pt(20)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    para = doc.add_paragraph()
    run = para.add_run('\n')
    run = para.add_run('파이썬교육기관장') 
    run.font.name = '나눔고딕'
    run.bold = True
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '나눔고딕')
    run.font.size = docx.shared.Pt(20)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.save(rf'/Users/devwonny/Documents/40products/12. 엑셀의 정보를 불러와 수료증 자동 생성/수료증결과{name_list[i]}.docx')
    convert(f'/Users/devwonny/Documents/40products/12. 엑셀의 정보를 불러와 수료증 자동 생성/수료증결과{name_list[i]}.docx', '/Users/devwonny/Documents/40products/12. 엑셀의 정보를 불러와 수료증 자동 생성/수료증결과test.pdf')